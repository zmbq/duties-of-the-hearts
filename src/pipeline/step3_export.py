"""
Step 3: Export translations to DOCX format.

This script creates Word documents with translations in a table format:
- Right column: Translated text (modern Hebrew)
- Left column: Original text (medieval Hebrew) - optional
- Narrow column: Paragraph numbers

Supports exporting individual sections, chapters, or the entire book.
"""

from typing import Optional, List
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from src.models import init_db, Chapter, Section, Paragraph, Translation
from src.utils import config


class DocumentExporter:
    """Service for exporting translations to DOCX format."""
    
    def __init__(self, show_original: bool = True, prompt_name: str = "modern"):
        """
        Initialize the document exporter.
        
        Args:
            show_original: If True, show original medieval Hebrew text alongside translation
            prompt_name: Name of the translation prompt to export
        """
        self.show_original = show_original
        self.prompt_name = prompt_name
        self.doc = Document()
        
        # Set RTL (right-to-left) for the document
        self._set_document_rtl()
        
        # Configure default font for Hebrew
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'David'  # Hebrew font
        font.size = Pt(12)
    
    def _set_document_rtl(self):
        """Set the document to RTL (right-to-left) mode for Hebrew."""
        sections = self.doc.sections
        for section in sections:
            sectPr = section._sectPr
            bidi = sectPr.find(qn('w:bidi'))
            if bidi is None:
                bidi = OxmlElement('w:bidi')
                sectPr.append(bidi)
    
    def _fix_rtl_text(self, text: str) -> str:
        """
        Fix parentheses and brackets for RTL display.
        In RTL text, parentheses need to be reversed.
        
        Args:
            text: Text to fix
            
        Returns:
            Text with reversed parentheses and brackets
        """
        if not text:
            return text
        
        # Create translation table for RTL punctuation
        rtl_fixes = str.maketrans({
            '(': ')',
            ')': '(',
            '[': ']',
            ']': '[',
            '{': '}',
            '}': '{',
            '<': '>',
            '>': '<'
        })
        
        return text.translate(rtl_fixes)
    
    def _set_paragraph_rtl(self, paragraph):
        """Set a paragraph to RTL (right-to-left) mode."""
        pPr = paragraph._element.get_or_add_pPr()
        bidi = pPr.find(qn('w:bidi'))
        if bidi is None:
            bidi = OxmlElement('w:bidi')
            pPr.append(bidi)
    
    def _add_chapter_heading(self, chapter: Chapter):
        """Add a chapter heading to the document."""
        heading = self.doc.add_heading(level=1)
        run = heading.add_run(self._fix_rtl_text(chapter.title))
        run.font.name = 'David'
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        # Set RTL for Hebrew text
        self._set_paragraph_rtl(heading)
    
    def _add_section_heading(self, section: Section):
        """Add a section heading to the document."""
        heading = self.doc.add_heading(level=2)
        run = heading.add_run(self._fix_rtl_text(section.title))
        run.font.name = 'David'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(51, 102, 153)  # Medium blue
        # Set RTL for Hebrew text
        self._set_paragraph_rtl(heading)
    
    def _create_table_header(self):
        """Create table with header row."""
        if self.show_original:
            # 3 columns: Number | Original | Translation
            table = self.doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            
            # Set column widths
            table.columns[0].width = Inches(0.4)  # Narrow for numbers
            table.columns[1].width = Inches(3.0)  # Original text
            table.columns[2].width = Inches(3.0)  # Translation
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = '#'
            header_cells[1].text = self._fix_rtl_text('תרגום מקורי')
            header_cells[2].text = self._fix_rtl_text('עברית מודרנית')
            
            # Format header
            for cell in header_cells:
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._set_paragraph_rtl(para)  # Set RTL
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
                    run.font.name = 'David'
                # Shade header
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'D9E2F3')  # Light blue
                cell._element.get_or_add_tcPr().append(shading_elm)
        else:
            # 2 columns: Number | Translation
            table = self.doc.add_table(rows=1, cols=2)
            table.style = 'Light Grid Accent 1'
            
            # Set column widths
            table.columns[0].width = Inches(0.4)  # Narrow for numbers
            table.columns[1].width = Inches(6.0)  # Translation
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = '#'
            header_cells[1].text = self._fix_rtl_text('תרגום (עברית מודרנית)')
            
            # Format header
            for cell in header_cells:
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._set_paragraph_rtl(para)  # Set RTL
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
                    run.font.name = 'David'
                # Shade header
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'D9E2F3')  # Light blue
                cell._element.get_or_add_tcPr().append(shading_elm)
        
        return table
    
    def _add_paragraph_row(self, table, para: Paragraph, translation: Translation):
        """Add a row to the table with paragraph data."""
        row = table.add_row()
        cells = row.cells
        
        # Paragraph number
        cells[0].text = str(para.paragraph_number)
        num_para = cells[0].paragraphs[0]
        num_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in num_para.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(128, 128, 128)  # Gray
        
        if self.show_original:
            # Original text
            cells[1].text = self._fix_rtl_text(para.text)
            orig_para = cells[1].paragraphs[0]
            self._set_paragraph_rtl(orig_para)  # Set RTL
            for run in orig_para.runs:
                run.font.name = 'David'
                run.font.size = Pt(11)
            
            # Translation
            cells[2].text = self._fix_rtl_text(translation.translated_text)
            trans_para = cells[2].paragraphs[0]
            self._set_paragraph_rtl(trans_para)  # Set RTL
            for run in trans_para.runs:
                run.font.name = 'David'
                run.font.size = Pt(11)
        else:
            # Translation only
            cells[1].text = self._fix_rtl_text(translation.translated_text)
            trans_para = cells[1].paragraphs[0]
            self._set_paragraph_rtl(trans_para)  # Set RTL
            for run in trans_para.runs:
                run.font.name = 'David'
                run.font.size = Pt(11)
    
    def export_section(self, session, section: Section, include_heading: bool = True):
        """
        Export a single section to the document.
        
        Args:
            session: Database session
            section: Section to export
            include_heading: Whether to include chapter/section headings
        """
        if include_heading:
            self._add_chapter_heading(section.chapter)
            self._add_section_heading(section)
        
        # Get paragraphs and their translations
        paragraphs = session.query(Paragraph).filter(
            Paragraph.section_id == section.id
        ).order_by(Paragraph.paragraph_number).all()
        
        if not paragraphs:
            self.doc.add_paragraph("(אין פסקאות בסעיף זה)")
            return
        
        # Create table
        table = self._create_table_header()
        
        # Add rows
        for para in paragraphs:
            # Get translation
            translation = session.query(Translation).filter(
                Translation.paragraph_id == para.id,
                Translation.prompt_name == self.prompt_name
            ).first()
            
            if translation:
                self._add_paragraph_row(table, para, translation)
            else:
                # No translation yet
                row = table.add_row()
                cells = row.cells
                cells[0].text = str(para.paragraph_number)
                cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                if self.show_original:
                    cells[1].text = self._fix_rtl_text(para.text)
                    orig_para = cells[1].paragraphs[0]
                    self._set_paragraph_rtl(orig_para)
                    for run in orig_para.runs:
                        run.font.name = 'David'
                    
                    cells[2].text = self._fix_rtl_text("(טרם תורגם)")
                    untrans_para = cells[2].paragraphs[0]
                    self._set_paragraph_rtl(untrans_para)
                    for run in untrans_para.runs:
                        run.font.italic = True
                        run.font.color.rgb = RGBColor(128, 128, 128)
                        run.font.name = 'David'
                else:
                    cells[1].text = self._fix_rtl_text("(טרם תורגם)")
                    untrans_para = cells[1].paragraphs[0]
                    self._set_paragraph_rtl(untrans_para)
                    for run in untrans_para.runs:
                        run.font.italic = True
                        run.font.color.rgb = RGBColor(128, 128, 128)
                        run.font.name = 'David'
        
        # Add spacing after table
        self.doc.add_paragraph()
    
    def export_chapter(self, session, chapter: Chapter):
        """
        Export an entire chapter (all sections) to the document.
        
        Args:
            session: Database session
            chapter: Chapter to export
        """
        self._add_chapter_heading(chapter)
        
        # Check if chapter has sections
        sections = session.query(Section).filter(
            Section.chapter_id == chapter.id
        ).order_by(Section.section_number).all()
        
        if sections:
            # Chapter has sections
            for section in sections:
                self._add_section_heading(section)
                
                # Export section content (without headings since we just added them)
                paragraphs = session.query(Paragraph).filter(
                    Paragraph.section_id == section.id
                ).order_by(Paragraph.paragraph_number).all()
                
                if paragraphs:
                    table = self._create_table_header()
                    for para in paragraphs:
                        translation = session.query(Translation).filter(
                            Translation.paragraph_id == para.id,
                            Translation.prompt_name == self.prompt_name
                        ).first()
                        
                        if translation:
                            self._add_paragraph_row(table, para, translation)
                    
                    self.doc.add_paragraph()
        else:
            # Chapter has no sections - direct paragraphs
            paragraphs = session.query(Paragraph).filter(
                Paragraph.chapter_id == chapter.id,
                Paragraph.section_id == None
            ).order_by(Paragraph.paragraph_number).all()
            
            if paragraphs:
                table = self._create_table_header()
                for para in paragraphs:
                    translation = session.query(Translation).filter(
                        Translation.paragraph_id == para.id,
                        Translation.prompt_name == self.prompt_name
                    ).first()
                    
                    if translation:
                        self._add_paragraph_row(table, para, translation)
                
                self.doc.add_paragraph()
    
    def save(self, filename: Path):
        """Save the document to a file."""
        self.doc.save(filename)
        print(f"✓ Document saved: {filename}")


def main():
    """Main function - export the test section."""
    import argparse
    
    parser = argparse.ArgumentParser(
        description='Export translations to DOCX format'
    )
    parser.add_argument(
        '--prompt',
        type=str,
        default='modern',
        help='Prompt name to export (default: modern)'
    )
    parser.add_argument(
        '--chapter',
        type=int,
        help='Chapter number to export (default: 6 for testing)'
    )
    parser.add_argument(
        '--section',
        type=int,
        help='Section number to export (optional - if omitted, exports entire chapter)'
    )
    parser.add_argument(
        '--show-original',
        action='store_true',
        default=True,
        help='Show original text alongside translation (default: True)'
    )
    parser.add_argument(
        '--no-original',
        action='store_true',
        help='Hide original text, show translation only'
    )
    parser.add_argument(
        '--output',
        type=str,
        help='Output filename (default: auto-generated based on prompt and content)'
    )
    
    args = parser.parse_args()
    
    print("Starting document export...")
    
    # Initialize database
    engine, SessionLocal = init_db(config.database_url)
    session = SessionLocal()
    
    try:
        # Find the chapter
        chapter_num = args.chapter or 6
        
        chapter = session.query(Chapter).filter(
            Chapter.chapter_number == chapter_num
        ).first()
        
        if not chapter:
            print(f"Error: Could not find chapter {chapter_num}")
            return
        
        # Determine if we're exporting a section or entire chapter
        if args.section:
            # Export specific section
            section = session.query(Section).filter(
                Section.chapter_id == chapter.id,
                Section.section_number == args.section
            ).first()
            
            if not section:
                print(f"Error: Could not find section {args.section} in chapter {chapter_num}")
                return
            
            print(f"Exporting: {chapter.title} → {section.title}")
            print(f"Using prompt: {args.prompt}")
            
            # Determine if we should show original
            show_original = args.show_original and not args.no_original
            
            # Create output filename
            output_dir = config.output_dir
            if args.output:
                output_file = output_dir / args.output
            else:
                # Auto-generate filename
                original_suffix = "_with_original" if show_original else "_only"
                output_file = output_dir / f"ch{chapter_num}_sec{args.section}_{args.prompt}{original_suffix}.docx"
            
            # Create exporter
            print(f"\nCreating document (show_original={show_original})...")
            exporter = DocumentExporter(show_original=show_original, prompt_name=args.prompt)
            exporter.export_section(session, section)
        else:
            # Export entire chapter
            print(f"Exporting entire chapter: {chapter.title}")
            print(f"Using prompt: {args.prompt}")
            
            # Determine if we should show original
            show_original = args.show_original and not args.no_original
            
            # Create output filename
            output_dir = config.output_dir
            if args.output:
                output_file = output_dir / args.output
            else:
                # Auto-generate filename
                original_suffix = "_with_original" if show_original else "_only"
                output_file = output_dir / f"chapter{chapter_num}_{args.prompt}{original_suffix}.docx"
            
            # Create exporter
            print(f"\nCreating document (show_original={show_original})...")
            exporter = DocumentExporter(show_original=show_original, prompt_name=args.prompt)
            exporter.export_chapter(session, chapter)
        
        exporter.save(output_file)
        
        print("\n" + "=" * 80)
        print("Export Summary")
        print("=" * 80)
        print(f"✓ Document created: {output_file}")
        print(f"  Prompt: {args.prompt}")
        print(f"  Show original: {show_original}")
        print("\nOpen the file in Word to review the translation!")
        
    finally:
        session.close()


if __name__ == "__main__":
    main()
